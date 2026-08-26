"""Gera a versao .docx do diario de desenvolvimento (mesmo conteudo do
artifact HTML), para uso offline sem depender de link externo."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

OUT = Path(r"c:\Automacao\Warranty Report Builder - Diario de Desenvolvimento.docx")

INK = RGBColor(0x1E, 0x24, 0x20)
MUTED = RGBColor(0x5B, 0x65, 0x60)
ACCENT = RGBColor(0xB2, 0x57, 0x1A)
ACCENT2 = RGBColor(0x2E, 0x6E, 0x62)
CODE_BG = "1B2420"
CODE_INK = RGBColor(0xDC, 0xE7, 0xE0)
CALLOUT_BG = "F2E3D6"
FORMULA_BG = "F4F5F1"
TABLE_HEAD_BG = "E1EEEA"

MONO = "Consolas"
BODY_FONT = "Georgia"


def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def border_paragraph(paragraph, color="D8DBD3", left_color=None, left_size=24):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(left_size if side == "left" and left_color else 6))
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), left_color if (side == "left" and left_color) else color)
        pBdr.append(el)
    pPr.append(pBdr)


def set_run_inline_code(run):
    run.font.name = MONO
    run.font.size = Pt(9.5)
    run.font.color.rgb = ACCENT2
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), TABLE_HEAD_BG)
    rPr.append(shd)


def add_mixed_paragraph(doc, parts, space_after=10):
    """parts: lista de (texto, is_code, is_bold)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for part in parts:
        text, is_code, is_bold = (part + (False,) * 3)[:3]
        run = p.add_run(text)
        if is_code:
            set_run_inline_code(run)
        else:
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            run.font.color.rgb = INK
            run.bold = is_bold
    return p


def add_p(doc, text, space_after=10, italic=False, color=None, size=11, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    run.font.color.rgb = color or INK
    return p


def add_title(doc, eyebrow, title, lede, tags):
    add_p(doc, eyebrow.upper(), color=ACCENT, size=10, space_after=4, bold=True)
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.font.name = MONO
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = INK
    h.paragraph_format.space_after = Pt(10)
    add_p(doc, lede, color=MUTED, size=13, space_after=10)
    tagp = doc.add_paragraph()
    tagp.paragraph_format.space_after = Pt(24)
    for i, tag in enumerate(tags):
        r = tagp.add_run(("   " if i else "") + f"[{tag}]")
        r.font.name = MONO
        r.font.size = Pt(9)
        r.font.color.rgb = ACCENT2


def add_h1(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if number:
        r = p.add_run(f"{number}  ")
        r.font.name = MONO
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(15.5)
    r.bold = True
    r.font.color.rgb = INK
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "DCDFD7")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = ACCENT2
    return p


def add_sub(doc, text):
    add_p(doc, text, color=MUTED, size=10.5, italic=True, space_after=14)


def add_bullets(doc, items, numbered=False):
    for item in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.space_after = Pt(4)
        _fill_rich(p, item)


def _fill_rich(paragraph, item):
    """item: string simples OU lista de (texto, is_code, is_bold)"""
    if isinstance(item, str):
        r = paragraph.add_run(item)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = INK
        return
    for text, is_code, is_bold in [(p + (False,) * 3)[:3] for p in item]:
        r = paragraph.add_run(text)
        if is_code:
            set_run_inline_code(r)
        else:
            r.font.name = BODY_FONT
            r.font.size = Pt(11)
            r.font.color.rgb = INK
            r.bold = is_bold


def add_code(doc, code, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(2)
        r = lp.add_run(label)
        r.font.name = MONO
        r.font.size = Pt(8.5)
        r.font.color.rgb = MUTED
        r.italic = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.right_indent = Cm(0.25)
    shade_paragraph(p, CODE_BG)
    border_paragraph(p, color=CODE_BG)
    lines = code.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if line else " ")
        run.font.name = MONO
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_INK
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_callout(doc, label, text):
    lp = doc.add_paragraph()
    lp.paragraph_format.space_after = Pt(2)
    shade_paragraph(lp, CALLOUT_BG)
    border_paragraph(lp, left_color="B2571A", left_size=18)
    r = lp.add_run(label.upper())
    r.font.name = MONO
    r.font.size = Pt(8.5)
    r.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    shade_paragraph(p, CALLOUT_BG)
    border_paragraph(p, left_color="B2571A", left_size=18)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK


def add_formula(doc, intro_lines, eq_lines, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    shade_paragraph(p, FORMULA_BG)
    border_paragraph(p)
    for i, (text, bold) in enumerate(intro_lines):
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(10.5)
        r.bold = bold
        r.font.color.rgb = INK
        if i < len(intro_lines) - 1:
            r.add_break()

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2 if note else 14)
    shade_paragraph(p2, FORMULA_BG)
    border_paragraph(p2)
    for i, line in enumerate(eq_lines):
        r = p2.add_run(line)
        r.font.name = MONO
        r.font.size = Pt(10)
        r.font.color.rgb = ACCENT2
        if i < len(eq_lines) - 1:
            r.add_break()

    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(14)
        shade_paragraph(p3, FORMULA_BG)
        border_paragraph(p3)
        r = p3.add_run(note)
        r.font.name = BODY_FONT
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h.upper())
        r.font.name = MONO
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = MUTED
        shade_cell(hdr[i], TABLE_HEAD_BG)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            is_code = isinstance(val, tuple)
            text = val[0] if is_code else val
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if is_code:
                r.font.name = MONO
                r.font.color.rgb = ACCENT2
            else:
                r.font.name = BODY_FONT
                r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_diagram_text(doc, title, lines):
    lp = doc.add_paragraph()
    lp.paragraph_format.space_after = Pt(4)
    r = lp.add_run(title)
    r.font.name = MONO
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = MUTED
    add_code(doc, "\n".join(lines))


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    border = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "DCDFD7")
    border.append(top)
    p._p.get_or_add_pPr().append(border)


# =====================================================================

doc = Document()

section = doc.sections[0]
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)

style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(11)
style.font.color.rgb = INK

add_title(
    doc,
    "Diario de desenvolvimento - estudo de caso",
    "Warranty Report Builder",
    "Como uma planilha de garantia virou um executavel de duplo clique: leitura de Excel, "
    "montagem programatica de um .docx formatado e empacotamento em .exe sem dependencias "
    "- explicado etapa por etapa, arquivo por arquivo.",
    ["Python 3.14", "openpyxl", "python-docx", "tkinter", "PyInstaller", "OOXML / zip"],
)

add_h1(doc, "", "Visao geral do problema")
add_mixed_paragraph(doc, [
    ("O ponto de partida eram dois arquivos manuais: uma planilha ", False),
    ("Warranty 2026.xlsx", True),
    (" com os chamados de garantia do mes, e um ", False),
    ("Warranty Report.docx", True),
    (" formatado a mao - marca d'agua, cores da marca, uma tabela por chamado - que alguem "
     "preenchia copiando celula por celula. O objetivo era eliminar essa transcricao manual "
     "sem perder a formatacao, e entregar o resultado como um programa que qualquer pessoa "
     "da equipe pudesse abrir clicando duas vezes, sem instalar Python nem pedir permissao "
     "de administrador.", False),
])
add_p(doc, "Isso definiu tres exigencias que moldaram todas as decisoes tecnicas a seguir:")
add_bullets(doc, [
    [("Ler a planilha pelo ", False, False), ("nome", False, True), (" das colunas, nao pela posicao, para tolerar pequenas mudancas de layout.", False, False)],
    "Gerar o .docx clonando a formatacao ja existente em vez de recria-la do zero.",
    "Empacotar tudo - Python, bibliotecas e o modelo do relatorio - em um unico .exe portatil.",
])

add_h1(doc, "", "Arquitetura e arquivos")
add_sub(doc, "O projeto inteiro vive em c:\\Automacao\\warranty_app. Cada arquivo tem uma responsabilidade unica.")
add_diagram_text(doc, "estrutura de pastas", [
    "warranty_app/",
    "|-- app.py                    # interface grafica (tkinter) - o que o usuario roda",
    "|-- report_engine.py          # motor: le o xlsx, monta o docx",
    "|-- assets/",
    "|   `-- Warranty Report_template.docx   # modelo-semente empacotado no .exe",
    "|-- _tcl_data/, _tk_data/     # runtime do tkinter extraido (etapa 5)",
    "|-- build_seed_template.py    # script usado 1x para gerar o modelo-semente",
    "|-- strip_photos.py           # script usado 1x para remover fotos do docx original",
    "|-- fix_header_image.py       # script usado 1x para corrigir o cabecalho",
    "|-- extract_tcltk.py          # script usado 1x para extrair o Tcl/Tk",
    "`-- dist/",
    "    `-- WarrantyReportBuilder.exe        # entregavel final",
])
add_mixed_paragraph(doc, [
    ("Vale notar a diferenca entre os dois grupos de scripts: ", False, False),
    ("app.py", True), (" e ", False, False), ("report_engine.py", True),
    (" rodam toda vez que alguem usa o programa - sao o produto. Os outros quatro ("
     "build_seed_template.py, strip_photos.py, fix_header_image.py, extract_tcltk.py) sao "
     "ferramentas de bastidor, rodadas uma unica vez para preparar os arquivos que ficam "
     "dentro de assets/ e _tcl_data/ / _tk_data/. O usuario final nunca os executa.", False, False),
])
add_diagram_text(doc, "fluxo geral (planilha -> motor -> docx -> download)", [
    "  Warranty 2026.xlsx --le colunas (openpyxl)--\\",
    "                                                 >  report_engine.py  --doc.save()--> Warranty Report.docx",
    "  Warranty Report_template.docx (1 bloco-semente)/        (clona o bloco-semente         (pasta temporaria)",
    "                                                    1x por linha da planilha)                   |",
    "                                                                                    'Fazer download' (shutil.copy)",
    "                                                                                                   v",
    "                                                                                    arquivo final escolhido pelo usuario",
])

# ---------------- Etapa 1 ----------------
add_h1(doc, "01", "Investigacao dos arquivos originais")
add_sub(doc, "Antes de escrever qualquer logica, foi preciso entender a estrutura real dos dois arquivos de entrada.")

add_h2(doc, "A planilha")
add_mixed_paragraph(doc, [
    ("Abrindo ", False, False), ("Warranty 2026.xlsx", True),
    (" com ", False, False), ("openpyxl", True),
    (", apareceu uma armadilha: a coluna ", False, False), ("ID CHAMADO", True),
    (" guardava numeros como ", False, False), ("15.925", True),
    (" em vez de ", False, False), ("15925", True), (".", False, False),
])
add_callout(
    doc, "Descoberta",
    "O sistema de origem grava o numero com separador de milhar (15.925), mas o Excel/openpyxl "
    "interpreta o ponto como separador decimal - o valor real armazenado na celula e o float "
    "15.925, nao o inteiro 15925. Isso so foi percebido comparando a sequencia de protocolos ja "
    "existente no relatorio antigo (...15901, 15911) com os valores crus lidos da planilha.",
)

add_h2(doc, "O relatorio .docx")
add_mixed_paragraph(doc, [
    ("Um documento .docx e, por baixo dos panos, um arquivo ", False, False),
    (".zip", True), (" contendo XML (o padrao e chamado ", False, False),
    ("OOXML", False, True), ("). Para entender o modelo, ele foi extraido como zip e inspecionado diretamente:", False, False),
])
add_code(doc, (
    "import docx\n"
    "d = docx.Document('Warranty Report.docx')\n"
    "for t in d.tables:\n"
    "    print([c.text for c in t.rows[0].cells])"
), label="investigacao - python-docx")
add_p(doc, (
    "Isso revelou que cada chamado no relatorio e, na verdade, duas tabelas seguidas: uma com 7 "
    "linhas (Protocol, Date, Customer, Description, DOT, Site of Injury, Defects Description) e "
    "outra logo abaixo com o cabecalho \"DOT / INTERNAL / OCCURRENCE\" - essa segunda tabela "
    "cresce conforme a quantidade de pneus do chamado, repetindo um par de linhas (\"Tire N\" + "
    "linha em branco) para cada pneu extra. Entender essa regra foi a base de toda a Etapa 2."
))

# ---------------- Etapa 2 ----------------
add_h1(doc, "02", "O motor de geracao")
add_sub(doc, "Arquivo: report_engine.py - a unica parte do programa que sabe ler uma planilha e escrever um .docx.")

add_h2(doc, "Lendo a planilha por nome de coluna")
add_p(doc, (
    "Em vez de assumir \"a coluna B e o ID\", o codigo le a primeira linha da planilha e monta um "
    "dicionario nome -> posicao. Isso torna o programa tolerante a colunas reordenadas ou extras:"
))
add_code(doc, (
    "def read_warranty_rows(xlsx_path):\n"
    "    wb = openpyxl.load_workbook(xlsx_path, data_only=True)\n"
    "    ws = wb.active\n"
    "    header = [str(c.value).strip() if c.value is not None else \"\" for c in ws[1]]\n"
    "    col_index = {name: i for i, name in enumerate(header)}\n"
    "    missing = [c for c in REQUIRED_COLUMNS if c not in col_index]\n"
    "    if missing:\n"
    "        raise ValueError(\"Colunas ausentes na planilha: \" + \", \".join(missing))"
), label="report_engine.py - openpyxl")
add_p(doc, (
    "data_only=True pede ao openpyxl os valores calculados das celulas (nao as formulas cruas). "
    "wb.active pega a primeira planilha do arquivo. O restante da funcao percorre as linhas a "
    "partir da segunda (min_row=2) e ignora linhas totalmente vazias."
))

add_h2(doc, "A formula do protocolo")
add_p(doc, "Essa e a peca mais sutil do programa - resolve a armadilha encontrada na Etapa 1.")
add_formula(
    doc,
    [("Problema: a celula guarda 15.925 (float), mas o protocolo real e 15925.", False)],
    [
        "se o valor tem casas decimais -> protocolo = round(valor x 1000)",
        "15.925 x 1000 = 15925.0  ->  round(...) = 15925  [correto]",
    ],
    note=(
        "Por que funciona sempre: o separador de milhar do sistema de origem so entra a cada 3 "
        "digitos a partir da direita. Para os numeros desse relatorio (5 digitos, uma unica casa "
        "de milhar), isso produz sempre exatamente 3 casas decimais - entao multiplicar por 1000 "
        "desfaz o erro de interpretacao de forma exata, sem depender de \"chutar\" quantos digitos "
        "o protocolo tem."
    ),
)
add_code(doc, (
    "def format_protocol(raw_id):\n"
    "    if isinstance(raw_id, float):\n"
    "        if not raw_id.is_integer():\n"
    "            # a planilha grava \"15925\" como 15.925 (separador de milhar\n"
    "            # interpretado como ponto decimal ao abrir o arquivo)\n"
    "            return str(round(raw_id * 1000))\n"
    "        return str(int(raw_id))\n"
    "    if isinstance(raw_id, int):\n"
    "        return str(raw_id)\n"
    "    return str(raw_id).strip()"
))
add_p(doc, (
    "Note a guarda not raw_id.is_integer(): se um dia a planilha vier com o numero ja correto "
    "(ex.: 15925.0, sem parte fracionaria \"de verdade\"), a funcao nao aplica a correcao - ela "
    "so multiplica por 1000 quando detecta que o valor tem casas decimais genuinas, que e "
    "exatamente o sintoma do bug."
))

add_h2(doc, "Editando texto sem perder a formatacao")
add_p(doc, (
    "Uma celula de tabela no Word nao e so \"texto\" - o texto vive dentro de runs (w:r), e cada "
    "run carrega sua propria fonte, negrito, cor etc. Trocar cell.text diretamente apaga essa "
    "formatacao. A solucao foi escrever apenas no primeiro run e remover os demais:"
))
add_code(doc, (
    "def _set_cell_text(cell, text):\n"
    "    paragraphs = cell.paragraphs\n"
    "    first = paragraphs[0]\n"
    "    runs = first.runs\n"
    "    if runs:\n"
    "        runs[0].text = text\n"
    "        for extra in runs[1:]:\n"
    "            extra._element.getparent().remove(extra._element)\n"
    "    else:\n"
    "        first.add_run(text)\n"
    "    for extra_p in paragraphs[1:]:\n"
    "        extra_p._element.getparent().remove(extra_p._element)"
))
add_p(doc, (
    "Isso importa porque celulas como \"CUSTOMER\" no modelo original tinham o texto dividido em "
    "tres runs (codigo do cliente, uma tabulacao, o nome) - sobrescrever so o primeiro run e "
    "apagar o resto garante que a fonte/cor da celula permaneca igual a do modelo."
))

add_h2(doc, "Encontrando o modelo dentro do proprio documento")
add_p(doc, (
    "Em vez de guardar em codigo fixo \"a tabela de protocolo e a primeira tabela\", "
    "_find_templates procura pelo conteudo: a primeira tabela cuja celula (0,0) comeca com "
    "\"PROTOCOL\", a primeira linha cujo texto comeca com \"DOT\", a primeira linha \"Tire...\", "
    "etc. Isso torna a funcao resiliente a mudancas no modelo desde que os rotulos continuem os mesmos."
))
add_code(doc, (
    "for t in doc.tables:\n"
    "    label = t.rows[0].cells[0].text.strip().upper()\n"
    "    if label.startswith(\"PROTOCOL\") and table1_tbl is None:\n"
    "        table1_tbl = t._tbl\n"
    "    elif label.startswith(\"DOT\"):\n"
    "        if header_tr is None:\n"
    "            header_tr = t.rows[0]._tr\n"
    "        for row in t.rows[1:]:\n"
    "            first_cell_text = row.cells[0].text.strip().upper()\n"
    "            if first_cell_text.startswith(\"TIRE\") and label_tr is None:\n"
    "                label_tr = row._tr\n"
    "            elif first_cell_text == \"\" and blank_tr is None:\n"
    "                blank_tr = row._tr"
))
add_p(doc, (
    "Note o uso de t._tbl e row._tr: o python-docx e uma camada amigavel sobre XML puro (via "
    "lxml). Table e Row sao objetos Python que \"envolvem\" um elemento XML real (<w:tbl>, "
    "<w:tr>); o underscore expoe esse elemento cru sempre que e preciso manipular XML "
    "diretamente - coisa que a API \"bonita\" do python-docx nao oferece (como copiar uma linha "
    "inteira ou inserir um bloco em outro lugar do documento)."
))

add_h2(doc, "Montando a tabela de pneus dinamicamente")
add_p(doc, (
    "A quantidade de linhas dessa segunda tabela depende da coluna QUANTIDADE: 1 pneu = 2 linhas "
    "(cabecalho + 1 em branco), 2 pneus = 4 linhas (cabecalho + branco + \"Tire 2\" + branco), e "
    "assim por diante. _build_tires_table reconstroi essa tabela do zero a cada chamado:"
))
add_code(doc, (
    "new_tbl = deepcopy(source_tbl)\n"
    "for tr in list(new_tbl.findall(qn(\"w:tr\"))):\n"
    "    new_tbl.remove(tr)          # esvazia, mantem so a moldura da tabela\n\n"
    "new_tbl.append(deepcopy(header_tr_tmpl))\n"
    "for tire_index in range(1, quantity + 1):\n"
    "    if tire_index > 1:\n"
    "        new_tbl.append(deepcopy(label_tr_tmpl))   # linha \"Tire N\"\n"
    "    new_tbl.append(deepcopy(blank_tr_tmpl))       # linha de dados em branco"
))
add_p(doc, (
    "deepcopy (do modulo padrao copy) e essencial aqui: um elemento XML so pode ter um pai por "
    "vez. Sem copiar, tentar inserir a mesma linha-modelo duas vezes simplesmente moveria o "
    "mesmo no de um lugar para o outro, e nao duplicaria nada. qn(\"w:tr\") (de docx.oxml.ns) "
    "resolve o prefixo de namespace w: para a URI completa que o XML do Word realmente usa por "
    "baixo dos panos."
))

add_h2(doc, "O truque do bloco-semente: clonar e depois apagar o original")
add_p(doc, (
    "O modelo empacotado no programa nao e um documento \"em branco\" - ele contem um chamado de "
    "exemplo real, com toda a formatacao correta. A funcao abaixo clona esse bloco uma vez por "
    "linha da planilha e, no final, apaga o bloco original, deixando so os clones no resultado:"
))
add_diagram_text(doc, "estado do documento em tres momentos", [
    "modelo carregado:              [ SEMENTE ] [ sectPr ]",
    "                                      |",
    "                          insere N clones antes do sectPr",
    "                          (sectPr.addprevious)",
    "                                      v",
    "apos inserir clones:   [ SEMENTE ] [clone1] [clone2] ... [cloneN] [ sectPr ]",
    "                                      |",
    "                          remove os \"seed_elements\" originais",
    "                                      v",
    "resultado final:                 [clone1] [clone2] ... [cloneN] [ sectPr ]",
])
add_code(doc, (
    "def build_report_from_rows(rows):\n"
    "    doc = Document(bundled_template_path())\n"
    "    templates = _find_templates(doc)\n"
    "    body = doc.element.body\n"
    "    sectPr = body.find(qn(\"w:sectPr\"))\n"
    "    seed_elements = [el for el in body.iterchildren() if el is not sectPr]\n\n"
    "    for row in rows:\n"
    "        # ...monta heading_p, table1_el, table2_el a partir dos templates...\n"
    "        for el in (heading_p, spacer1, table1_el, spacer2, table2_el, spacer3):\n"
    "            sectPr.addprevious(el)     # insere sempre logo antes da secao final\n\n"
    "    for el in seed_elements:\n"
    "        body.remove(el)                # apaga o bloco-modelo original\n\n"
    "    return doc, added, duplicated"
))
add_p(doc, (
    "sectPr (section properties) e o elemento que fecha o corpo do documento e guarda pagina, "
    "margens e cabecalho/rodape - ele precisa continuar sendo o ultimo filho do corpo do "
    "documento. sectPr.addprevious(el) (metodo do lxml) insere el imediatamente antes dele; "
    "chamado varias vezes em sequencia, cada novo elemento entra \"colado\" antes do sectPr, "
    "preservando a ordem em que foram inseridos."
))

add_h2(doc, "O ponto de entrada")
add_p(doc, "Por fim, a funcao que a interface grafica chama:")
add_code(doc, (
    "def process_xlsx_into_report(xlsx_path):\n"
    "    rows = read_warranty_rows(xlsx_path)\n"
    "    doc, added, duplicated = build_report_from_rows(rows)\n"
    "    path = output_path()          # pasta temporaria do sistema (tempfile)\n"
    "    doc.save(path)\n"
    "    return added, duplicated, path"
))
add_p(doc, (
    "output_path() usa tempfile.gettempdir() - a pasta temporaria do Windows - em vez de salvar "
    "ao lado do executavel. Isso garante que o programa nunca precise de permissao de escrita em "
    "pastas protegidas (como Arquivos de Programas), cumprindo o requisito de rodar sem "
    "privilegios de administrador."
))

# ---------------- Etapa 3 ----------------
add_h1(doc, "03", "A interface grafica")
add_sub(doc, "Arquivo: app.py - uma janela tkinter fina, que so chama funcoes do report_engine.")
add_p(doc, (
    "tkinter e a biblioteca de interface grafica que ja vem instalada com o Python - nao precisa "
    "de nenhuma instalacao extra. A janela inteira e uma classe que herda de tk.Tk:"
))
add_code(doc, (
    "class App(tk.Tk):\n"
    "    def __init__(self):\n"
    "        super().__init__()\n"
    "        self.title(WINDOW_TITLE)\n"
    "        self.geometry(\"520x320\")\n"
    "        self.send_button = tk.Button(self, text=\"Enviar arquivo\", command=self.on_send_file)\n"
    "        self.download_button = tk.Button(self, text=\"Fazer download\", state=tk.DISABLED,\n"
    "                                          command=self.on_download)"
), label="app.py - tkinter")
add_p(doc, "Dois pontos merecem destaque:")
add_bullets(doc, [
    [("Botoes chamam funcoes, nao o contrario.", False, True),
     (" O parametro command=self.on_send_file registra um callback: o tkinter guarda essa "
      "referencia e so executa a funcao quando o clique realmente acontece. Nao existe um loop "
      "escrito a mao esperando o clique - App().mainloop(), na ultima linha do arquivo, e quem "
      "fica escutando eventos do sistema operacional indefinidamente.", False, False)],
    [("Estado simples guardado na propria instancia.", False, True),
     (" self.last_report_path guarda o caminho do ultimo relatorio gerado; o botao \"Fazer "
      "download\" comeca desabilitado e so e liberado depois de um processamento bem-sucedido - "
      "evitando que o usuario tente baixar antes de existir algo para baixar.", False, False)],
])

add_h2(doc, "As caixas de dialogo nativas")
add_code(doc, (
    "xlsx_path = filedialog.askopenfilename(\n"
    "    title=\"Selecione a planilha de garantia\",\n"
    "    filetypes=[(\"Planilha Excel\", \"*.xlsx\")],\n"
    ")\n"
    "...\n"
    "destino = filedialog.asksaveasfilename(\n"
    "    title=\"Salvar relatorio como\",\n"
    "    defaultextension=\".docx\",\n"
    "    initialfile=\"Warranty Report.docx\",\n"
    ")"
))
add_p(doc, (
    "tkinter.filedialog abre as janelas nativas do Windows para escolher/salvar arquivo - o "
    "mesmo dialogo que qualquer outro programa usa -, o que da ao aplicativo uma aparencia "
    "familiar sem esforco. messagebox.showerror cumpre o mesmo papel para avisos de erro. Todo o "
    "tratamento de erro no on_send_file e um try/except Exception ao redor da chamada ao motor - "
    "qualquer problema (planilha com coluna faltando, arquivo corrompido) vira uma mensagem "
    "legivel em vez de travar o programa."
))

# ---------------- Etapa 4 ----------------
add_h1(doc, "04", "Empacotando em um unico .exe")
add_sub(doc, "Ferramenta: PyInstaller - le o grafo de importacoes do programa e empacota tudo em um executavel.")
add_code(doc, (
    "python -m PyInstaller --noconfirm --onefile --windowed --name \"WarrantyReportBuilder\" `\n"
    "    --add-data \"assets;assets\" `\n"
    "    --add-data \"_tcl_data;_tcl_data\" `\n"
    "    --add-data \"_tk_data;_tk_data\" `\n"
    "    app.py"
), label="powershell - build final")
add_table(doc, ["Flag", "O que faz"], [
    ["--onefile", "Empacota tudo (interpretador + bibliotecas + dados) em um unico arquivo .exe."],
    ["--windowed", "Nao abre um console de terminal atras da janela do tkinter."],
    ["--name", "Define o nome do executavel final."],
    ["--add-data \"origem;destino\"", "Copia uma pasta/arquivo para dentro do pacote. No Windows o separador e ; (em Linux/Mac seria :)."],
])
add_p(doc, (
    "Em tempo de execucao, o .exe se extrai para uma pasta temporaria e disponibiliza o caminho "
    "dela na variavel sys._MEIPASS. E por isso que report_engine.py resolve o caminho do modelo assim:"
))
add_code(doc, (
    "def bundled_template_path():\n"
    "    base = Path(getattr(sys, \"_MEIPASS\", app_dir()))\n"
    "    return base / \"assets\" / TEMPLATE_FILENAME"
))
add_p(doc, (
    "getattr(sys, \"_MEIPASS\", app_dir()) le o atributo _MEIPASS se ele existir (dentro do .exe "
    "empacotado) ou cai de volta para a pasta do proprio script (rodando via python app.py "
    "normalmente, em desenvolvimento). O mesmo arquivo de codigo funciona nos dois cenarios sem "
    "nenhum \"if e exe\" espalhado pelo programa."
))

# ---------------- Etapa 5 ----------------
add_h1(doc, "05", "O obstaculo do Tcl/Tk")
add_sub(doc, "O primeiro .exe gerado abria e travava imediatamente. Esta foi a etapa mais trabalhosa do projeto.")
add_p(doc, (
    "O tkinter nao e Python puro - e uma casca fina sobre duas bibliotecas escritas em C, Tcl e "
    "Tk, que dependem de um conjunto de arquivos de script (.tcl) para funcionar. Normalmente o "
    "PyInstaller sabe localizar e empacotar esses arquivos automaticamente. So que, nesta "
    "instalacao especifica do Python, eles nao existiam como arquivos soltos no disco - estavam "
    "guardados dentro de um sistema de arquivos virtual compactado (zipfs) embutido na propria "
    "DLL do Tcl 9. Confirmado rodando isto:"
))
add_code(doc, (
    "r = tkinter.Tk()\n"
    "print(r.tk.eval('set tcl_library'))\n"
    "# //zipfs:/lib/tcl/tcl_library   <- nao e um caminho de disco de verdade"
))
add_p(doc, (
    "Como o PyInstaller so sabe copiar arquivos reais, e esses arquivos \"moravam\" dentro de um "
    "zip embutido na DLL, a solucao foi pedir para o proprio interpretador Tcl copiar seus "
    "arquivos internos para uma pasta real, usando os comandos nativos do Tcl (glob, file copy) "
    "atraves do tkinter:"
))
add_code(doc, (
    "r = tkinter.Tk()\n"
    "script = r\"\"\"\n"
    "proc copyDir {src dst} {\n"
    "    file mkdir $dst\n"
    "    foreach f [glob -nocomplain -directory $src *] {\n"
    "        set name [file tail $f]\n"
    "        set target [file join $dst $name]\n"
    "        if {[file isdirectory $f]} { copyDir $f $target } else { file copy -force $f $target }\n"
    "    }\n"
    "}\n"
    "copyDir {//zipfs:/lib/tcl/tcl_library} {.../_tcl_data}\n"
    "copyDir {//zipfs:/lib/tk/tk_library} {.../_tk_data}\n"
    "\"\"\"\n"
    "r.tk.eval(script)   # executa o script Tcl acima dentro do proprio interpretador"
), label="extract_tcltk.py - script de bastidor, roda 1x")
add_p(doc, (
    "r.tk.eval(...) envia texto diretamente para o interpretador Tcl embutido no tkinter - e uma "
    "\"porta dos fundos\" para rodar comandos Tcl nativos a partir do Python. O resultado foram "
    "duas pastas reais (_tcl_data, _tk_data, 964 arquivos ao todo) prontas para serem "
    "empacotadas normalmente com --add-data."
))
add_callout(
    doc, "Detalhe que quase passou despercebido",
    "Nao basta empacotar essas pastas com qualquer nome - o proprio PyInstaller ja injeta, na "
    "inicializacao do .exe, um script (pyi_rth__tkinter.py) que procura por pastas chamadas "
    "exatamente _tcl_data e _tk_data na raiz do pacote e define TCL_LIBRARY/TK_LIBRARY "
    "automaticamente. A primeira tentativa usou nomes diferentes e continuou falhando - so "
    "funcionou depois de renomear as pastas para bater exatamente com o que esse script interno "
    "ja esperava.",
)

# ---------------- Etapa 6 ----------------
add_h1(doc, "06", "Reduzindo 168 MB para 12,5 MB")
add_sub(doc, "Arquivo: strip_photos.py - o relatorio original trazia 390 fotos de evidencia coladas no corpo do documento.")
add_p(doc, (
    "Como um .docx e um zip, da para inspeciona-lo com a biblioteca padrao zipfile, sem precisar "
    "do Word nem do python-docx. Isso revelou que as imagens nao eram \"lixo\" orfao - 388 delas "
    "estavam de fato referenciadas em word/_rels/document.xml.rels e inseridas inline no corpo "
    "(word/document.xml), uma por chamado."
))
add_code(doc, (
    "DRAWING_RE = re.compile(rb\"<w:drawing\\b.*?</w:drawing>\", re.DOTALL)\n"
    "IMAGE_REL_RE = re.compile(\n"
    "    rb'<Relationship [^>]*Type=\"[^\"]*?/relationships/image\"[^>]*Target=\"/media/[^\"]*\"[^>]*/>'\n"
    ")\n\n"
    "with zipfile.ZipFile(SRC, \"r\") as zin:\n"
    "    document_xml = DRAWING_RE.sub(b\"\", zin.read(\"word/document.xml\"))\n"
    "    rels_xml = IMAGE_REL_RE.sub(b\"\", zin.read(\"word/_rels/document.xml.rels\"))\n"
    "    with zipfile.ZipFile(DST, \"w\", zipfile.ZIP_DEFLATED) as zout:\n"
    "        for item in zin.infolist():\n"
    "            if item.filename == \"word/document.xml\":\n"
    "                zout.writestr(item, document_xml)\n"
    "            elif item.filename == \"word/_rels/document.xml.rels\":\n"
    "                zout.writestr(item, rels_xml)\n"
    "            elif item.filename.startswith(\"media/\"):\n"
    "                continue              # nao copia as fotos para o novo zip\n"
    "            else:\n"
    "                zout.writestr(item, zin.read(item.filename))"
), label="strip_photos.py - zipfile + re")
add_p(doc, "Tres operacoes coordenadas, cada uma resolvendo uma \"ponta\" da mesma referencia:")
add_bullets(doc, [
    [("Remover a tag <w:drawing>", False, True),
     (" do corpo do documento (a expressao regular usa re.DOTALL para que o . tambem combine "
      "quebras de linha, ja que o bloco de uma imagem pode ocupar varias linhas de XML).", False, False)],
    [("Remover a entrada correspondente", False, True),
     (" em document.xml.rels - o arquivo que traduz o \"nome interno\" usado no XML para o "
      "caminho real do arquivo de imagem dentro do zip.", False, False)],
    [("Nao copiar os arquivos de imagem", False, True),
     (" (pasta media/) para o novo zip.", False, False)],
], numbered=True)
add_p(doc, (
    "Pular qualquer uma das tres deixaria o Word reclamando de \"conteudo ilegivel\" (referencia "
    "quebrada) ou o arquivo do mesmo tamanho de antes (imagens continuariam la, so sem "
    "aparecer). Feitas as tres, o cabecalho com a marca d'agua - que e uma imagem separada, "
    "referenciada por um arquivo de cabecalho proprio (header1.xml) - nao foi tocado, porque o "
    "filtro busca so relacionamentos vindos de document.xml.rels."
))

# ---------------- Etapa 7 ----------------
add_h1(doc, "07", "O modelo-semente")
add_sub(doc, "Arquivo: build_seed_template.py - prepara o arquivo que fica dentro de assets/ e viaja empacotado no .exe.")
add_p(doc, (
    "A primeira versao do programa mantinha o historico completo de chamados dentro do modelo, "
    "so acrescentando os novos a cada envio. Isso trazia dois problemas: o arquivo ficava pesado "
    "para sempre, e cada relatorio gerado carregava chamados de meses anteriores que nao tinham "
    "nada a ver com a planilha enviada. A solucao foi reduzir o modelo a um unico bloco de "
    "exemplo - usado so como referencia de formatacao - que o motor descarta antes de salvar "
    "(ver a Etapa 2)."
))
add_code(doc, (
    "heading_indices = [idx for idx, el in enumerate(children) if \"Warranty Report\" in get_text(el)]\n"
    "cut_start = heading_indices[1]                       # inicio do 2o bloco\n"
    "cut_end = children.index(sect_pr)                    # onde a secao termina\n"
    "for el in children[cut_start:cut_end]:\n"
    "    body.remove(el)                                  # apaga do 2o bloco em diante"
))
add_p(doc, (
    "A logica localiza todos os paragrafos-titulo (\"Warranty Report...\") que abrem cada bloco "
    "de chamado, guarda o indice do segundo (ou seja, preserva o primeiro bloco inteiro) e "
    "remove tudo entre esse ponto e o sectPr. O resultado: um .docx de ~12,4 MB (a maior parte "
    "agora e a fonte da marca embutida, nao conteudo) com exatamente um chamado de exemplo."
))

# ---------------- Etapa 8 ----------------
add_h1(doc, "08", "Ajustes finos")
add_sub(doc, "Arquivo: fix_header_image.py - dois pedidos de refinamento depois de ver o resultado gerado.")

add_h2(doc, "Removendo a linha \"DATE\"")
add_p(doc, (
    "A primeira versao deixava a celula de data em branco; a versao final remove a linha inteira "
    "da tabela, direto em report_engine.py:"
))
add_code(doc, (
    "table1 = Table(table1_el, doc)\n"
    "date_row = next(r for r in table1.rows if r.cells[0].text.strip().upper() == \"DATE\")\n"
    "table1_el.remove(date_row._tr)     # remove o <w:tr> do XML da tabela\n"
    "table1 = Table(table1_el, doc)     # reconstroi o objeto: linhas seguintes agora tem novos indices"
))
add_p(doc, (
    "O ponto sutil e reconstruir o objeto Table depois de remover a linha: o python-docx calcula "
    "table.rows a partir do XML atual toda vez que a propriedade e acessada, entao recriar o "
    "wrapper garante que table1.rows[1] agora aponte para \"CUSTOMER\" em vez de para a linha "
    "\"DATE\" que acabou de ser removida."
))

add_h2(doc, "Cabecalho/rodape de borda a borda")
add_p(doc, (
    "A marca d'agua e uma unica imagem alta o suficiente para cobrir a pagina inteira (topo e "
    "base, com uma faixa preta desenhada em cada ponta da propria imagem - por isso "
    "\"cabecalho e rodape\" vem da mesma figura). Ela e posicionada por um elemento VML (o "
    "formato de desenho legado que o Word usa especificamente para marcas d'agua) dentro de "
    "word/header1.xml:"
))
add_formula(
    doc,
    [
        ("Antes - posicao relativa a margem da pagina:", True),
        ("margin-top: -72.4pt   /   mso-position-vertical-relative: margin", False),
        ("", False),
        ("Depois - posicao relativa a pagina inteira:", True),
        ("top: 0pt   /   mso-position-vertical-relative: page", False),
    ],
    [],
    note=(
        "Ancorar em 'margin' faz a posicao depender do valor exato da margem configurada na "
        "secao (nesse caso, 1 polegada) - qualquer pequena diferenca entre o deslocamento "
        "negativo previsto e a margem real do documento sobra como uma faixa em branco visivel "
        "antes da imagem. Ancorar direto em 'page' elimina essa dependencia: top:0pt sempre "
        "significa \"encoste na borda fisica da pagina\", nao importa como as margens do "
        "documento estejam configuradas."
    ),
)
add_code(doc, (
    "OLD = (\"position:absolute;left:0pt;margin-left:-73.15pt;margin-top:-72.4pt;\"\n"
    "       \"height:792pt;width:613.5pt;mso-position-horizontal-relative:margin;\"\n"
    "       \"mso-position-vertical-relative:margin;\")\n"
    "NEW = (\"position:absolute;left:0pt;top:0pt;margin-left:0pt;margin-top:0pt;\"\n"
    "       \"height:792pt;width:612pt;mso-position-horizontal-relative:page;\"\n"
    "       \"mso-position-vertical-relative:page;\")\n"
    "header_xml = header_xml.replace(OLD, NEW)"
))
add_p(doc, (
    "792pt x 612pt equivalem exatamente a 11 x 8,5 polegadas - o tamanho de uma pagina Carta "
    "(Letter) - garantindo que a imagem cubra a folha inteira, ponta a ponta, em vez de apenas "
    "se aproximar do tamanho da margem."
))

# ---------------- Referencia ----------------
add_h1(doc, "", "Referencia rapida de bibliotecas")
add_table(doc, ["Biblioteca / modulo", "Para que serviu aqui"], [
    [("openpyxl",), "Ler celulas e formatos de planilhas .xlsx (nao vem com o Python - instalada via pip)."],
    [("python-docx",), "API de alto nivel para documentos .docx (paragrafos, tabelas, runs). Por baixo, usa lxml."],
    [("docx.oxml.ns.qn",), "Traduz prefixos curtos de namespace XML (w:tr) para a URI completa exigida pelo lxml."],
    [("copy.deepcopy",), "Duplica um elemento XML antes de inserir em outro lugar - um mesmo no nao pode ter dois pais."],
    [("zipfile",), "Biblioteca padrao para ler/escrever .zip - usada porque um .docx e um zip por dentro."],
    [("re",), "Expressoes regulares, para localizar e remover blocos <w:drawing> e relacionamentos de imagem."],
    [("tkinter",), "Interface grafica - ja incluida no Python, sem instalacao extra."],
    [("filedialog / messagebox",), "Dialogos nativos do sistema operacional para abrir, salvar e avisar erros."],
    [("pathlib.Path",), "Manipulacao de caminhos de arquivo de forma legivel e independente de sistema operacional."],
    [("tempfile",), "Localiza a pasta temporaria do sistema - evita precisar de permissao de administrador."],
    [("shutil",), "Copia de arquivos (usada no botao \"Fazer download\")."],
    [("sys",), "Detecta se o codigo roda como script normal ou dentro do .exe (sys.frozen, sys._MEIPASS)."],
    ["PyInstaller", "Ferramenta de linha de comando que empacota Python + bibliotecas + dados em um executavel."],
])

# ---------------- Comandos ----------------
add_h1(doc, "", "Comandos usados, na ordem")
add_code(doc, "python -m pip install openpyxl python-docx\npython -m pip install pyinstaller", label="powershell - instalacao")
add_code(doc, "python extract_tcltk.py", label="powershell - extrair tcl/tk (etapa 5, 1x)")
add_code(doc, (
    "python -m PyInstaller --noconfirm --onefile --windowed --name \"WarrantyReportBuilder\" `\n"
    "  --add-data \"assets;assets\" `\n"
    "  --add-data \"_tcl_data;_tcl_data\" `\n"
    "  --add-data \"_tk_data;_tk_data\" `\n"
    "  app.py"
), label="powershell - gerar build final")
add_p(doc, (
    "O acento grave (`) no fim das linhas e o caractere de continuacao de linha do PowerShell - "
    "o comando e uma unica chamada quebrada em varias linhas so para facilitar a leitura."
))

add_divider(doc)
add_p(doc, (
    "Documento gerado a partir do codigo-fonte real do projeto em c:\\Automacao\\warranty_app - "
    "todos os trechos mostrados correspondem ao estado atual dos arquivos, sem simplificacoes inventadas."
), color=MUTED, size=9.5, italic=True)

doc.save(OUT)
print(f"Salvo em: {OUT}")

"""Logica de leitura da planilha de garantia e geracao do Warranty Report.docx.

Cada planilha enviada gera um relatorio NOVO, contendo somente os registros
daquele envio. O modelo empacotado traz um unico bloco de exemplo (usado so
como referencia de formatacao/marca d'agua para clonar os blocos novos) que e
removido do documento antes de salvar o resultado final.
"""

import sys
import tempfile
from copy import deepcopy
from pathlib import Path

import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

REQUIRED_COLUMNS = [
    "ID CHAMADO",
    "CODIGO CLIENTE",
    "CLIENTE ORIGEM",
    "ITEM",
    "QUANTIDADE",
    "SUBTIPO",
    "DOT",
]

REPORT_FILENAME = "Warranty Report.docx"
TEMPLATE_FILENAME = "Warranty Report_template.docx"


def app_dir():
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def bundled_template_path():
    base = Path(getattr(sys, "_MEIPASS", app_dir()))
    return base / "assets" / TEMPLATE_FILENAME


def output_path():
    return Path(tempfile.gettempdir()) / REPORT_FILENAME


def format_protocol(raw_id):
    if isinstance(raw_id, float):
        if not raw_id.is_integer():
            # a planilha grava "15925" como 15.925 (separador de milhar
            # interpretado como ponto decimal ao abrir o arquivo)
            return str(round(raw_id * 1000))
        return str(int(raw_id))
    if isinstance(raw_id, int):
        return str(raw_id)
    return str(raw_id).strip()


def read_warranty_rows(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    header = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]
    col_index = {name: i for i, name in enumerate(header)}
    missing = [c for c in REQUIRED_COLUMNS if c not in col_index]
    if missing:
        raise ValueError("Colunas ausentes na planilha: " + ", ".join(missing))

    rows = []
    for excel_row in ws.iter_rows(min_row=2):
        values = [c.value for c in excel_row]
        if all(v in (None, "") for v in values):
            continue
        raw_id = values[col_index["ID CHAMADO"]]
        if raw_id in (None, ""):
            continue
        rows.append({
            "protocol": format_protocol(raw_id),
            "codigo_cliente": str(values[col_index["CODIGO CLIENTE"]] or "").strip(),
            "cliente_origem": str(values[col_index["CLIENTE ORIGEM"]] or "").strip(),
            "item": str(values[col_index["ITEM"]] or "").strip(),
            "quantidade": int(values[col_index["QUANTIDADE"]] or 1),
            "subtipo": str(values[col_index["SUBTIPO"]] or "").strip(),
            "dot": str(values[col_index["DOT"]] or "").strip(),
        })
    return rows


def _set_cell_text(cell, text):
    paragraphs = cell.paragraphs
    first = paragraphs[0]
    runs = first.runs
    if runs:
        runs[0].text = text
        for extra in runs[1:]:
            extra._element.getparent().remove(extra._element)
    else:
        first.add_run(text)
    for extra_p in paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)


def _find_templates(doc):
    heading_p = None
    spacer_p = None
    for p in doc.paragraphs:
        text = p.text
        if heading_p is None and "Warranty Report" in text:
            heading_p = p._p
        elif heading_p is not None and spacer_p is None and text.strip() == "":
            spacer_p = p._p
        if heading_p is not None and spacer_p is not None:
            break

    table1_tbl = None
    header_tr = None
    blank_tr = None
    label_tr = None
    for t in doc.tables:
        label = t.rows[0].cells[0].text.strip().upper()
        if label.startswith("PROTOCOL") and table1_tbl is None:
            table1_tbl = t._tbl
        elif label.startswith("DOT"):
            if header_tr is None:
                header_tr = t.rows[0]._tr
            for row in t.rows[1:]:
                first_cell_text = row.cells[0].text.strip().upper()
                if first_cell_text.startswith("TIRE") and label_tr is None:
                    label_tr = row._tr
                elif first_cell_text == "" and blank_tr is None:
                    blank_tr = row._tr

    missing = [name for name, val in [
        ("titulo do bloco", heading_p),
        ("paragrafo em branco", spacer_p),
        ("tabela de protocolo", table1_tbl),
        ("cabecalho DOT/INTERNAL/OCCURRENCE", header_tr),
        ("linha em branco da tabela de pneus", blank_tr),
        ("linha 'Tire N' da tabela de pneus", label_tr),
    ] if val is None]
    if missing:
        raise ValueError(
            "Nao encontrei no Warranty Report.docx um modelo de: " + ", ".join(missing)
        )

    return {
        "heading_p": heading_p,
        "spacer_p": spacer_p,
        "table1": table1_tbl,
        "header_tr": header_tr,
        "blank_tr": blank_tr,
        "label_tr": label_tr,
    }


def _build_tires_table(doc, templates, quantity):
    header_tr_tmpl = templates["header_tr"]
    blank_tr_tmpl = templates["blank_tr"]
    label_tr_tmpl = templates["label_tr"]

    source_tbl = header_tr_tmpl.getparent()
    new_tbl = deepcopy(source_tbl)
    for tr in list(new_tbl.findall(qn("w:tr"))):
        new_tbl.remove(tr)

    new_tbl.append(deepcopy(header_tr_tmpl))
    label_rows = []
    for tire_index in range(1, quantity + 1):
        if tire_index > 1:
            label_tr = deepcopy(label_tr_tmpl)
            new_tbl.append(label_tr)
            label_rows.append((label_tr, tire_index))
        new_tbl.append(deepcopy(blank_tr_tmpl))

    table = Table(new_tbl, doc)
    for label_tr, tire_index in label_rows:
        row = next(r for r in table.rows if r._tr is label_tr)
        for cell in row.cells:
            _set_cell_text(cell, f"Tire {tire_index}")
    label_tr_set = {id(tr) for tr, _ in label_rows}
    for row in table.rows[1:]:
        if id(row._tr) not in label_tr_set:
            for cell in row.cells:
                _set_cell_text(cell, "")

    return new_tbl


def build_report_from_rows(rows):
    """Monta um Warranty Report.docx novo contendo somente os `rows` informados.

    Parte do modelo-semente empacotado (1 bloco de exemplo, usado so como
    referencia de formatacao/marca d'agua) e remove esse bloco de exemplo do
    resultado final. Retorna (doc, added, duplicated).
    """
    doc = Document(bundled_template_path())
    templates = _find_templates(doc)
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    seed_elements = [el for el in body.iterchildren() if el is not sectPr]

    seen = set()
    added, duplicated = [], []
    for row in rows:
        if row["protocol"] in seen:
            duplicated.append(row["protocol"])
            continue
        seen.add(row["protocol"])

        heading_p = deepcopy(templates["heading_p"])
        spacer1 = deepcopy(templates["spacer_p"])
        table1_el = deepcopy(templates["table1"])
        spacer2 = deepcopy(templates["spacer_p"])
        table2_el = _build_tires_table(doc, templates, max(row["quantidade"], 1))
        spacer3 = deepcopy(templates["spacer_p"])

        table1 = Table(table1_el, doc)
        date_row = next(
            r for r in table1.rows if r.cells[0].text.strip().upper() == "DATE"
        )
        table1_el.remove(date_row._tr)
        table1 = Table(table1_el, doc)

        _set_cell_text(table1.rows[0].cells[1], row["protocol"])
        customer = f'{row["codigo_cliente"]} {row["cliente_origem"]}'.strip()
        _set_cell_text(table1.rows[1].cells[1], customer)
        _set_cell_text(table1.rows[2].cells[1], row["item"])
        _set_cell_text(table1.rows[3].cells[1], row["dot"])
        _set_cell_text(table1.rows[4].cells[1], row["subtipo"])
        _set_cell_text(table1.rows[5].cells[1], row["subtipo"])

        for el in (heading_p, spacer1, table1_el, spacer2, table2_el, spacer3):
            sectPr.addprevious(el)

        added.append(row["protocol"])

    for el in seed_elements:
        body.remove(el)

    return doc, added, duplicated


def process_xlsx_into_report(xlsx_path):
    """Le a planilha e gera um Warranty Report.docx novo (arquivo temporario).

    Retorna (added, duplicated, output_file_path).
    """
    rows = read_warranty_rows(xlsx_path)
    doc, added, duplicated = build_report_from_rows(rows)
    path = output_path()
    doc.save(path)
    return added, duplicated, path
